import os
import sys
import traceback
from PyPDF2 import PdfWriter, PdfReader
import pypdfium2 as pdfium


def reconstruir_layout_ocr_estruturado(res, img_width, img_height, printable_width_in=7.27):
    """
    Analisa o resultado do OCR e agrupa as palavras em linhas e blocos horizontais (clusters),
    retornando as posições exatas em polegadas para aplicação de Tab Stops no Word.
    """
    if not res or not res.get("lines"):
        return []

    words = []
    for line in res.get("lines", []):
        for w in line.get("words", []):
            rect = w.get("bounding_rect", {})
            text = w.get("text", "").strip()
            if text and rect:
                words.append({
                    "text": text,
                    "x": rect.get("x", 0),
                    "y": rect.get("y", 0),
                    "w": rect.get("width", 0),
                    "h": rect.get("height", 0)
                })

    if not words:
        return []

    # Ordenar palavras pelo eixo Y para agrupar em linhas
    words_sorted = sorted(words, key=lambda wd: wd["y"])
    linhas = []
    current_line = []
    tolerance_y = 16  # Tolerância vertical de pixels no mesmo nível de linha

    for w in words_sorted:
        if not current_line:
            current_line.append(w)
        else:
            avg_y = sum(x["y"] for x in current_line) / len(current_line)
            if abs(w["y"] - avg_y) <= tolerance_y:
                current_line.append(w)
            else:
                linhas.append(current_line)
                current_line = [w]
    if current_line:
        linhas.append(current_line)

    resultado_linhas = []

    for line_words in linhas:
        # Ordenar palavras da linha pelo eixo X
        line_words.sort(key=lambda wd: wd["x"])
        
        # Agrupar em clusters na mesma linha baseando-se na distância horizontal
        clusters = []
        current_cluster = []
        
        for w in line_words:
            if not current_cluster:
                current_cluster.append(w)
            else:
                last_w = current_cluster[-1]
                gap = w["x"] - (last_w["x"] + last_w["w"])
                # Se a distância entre palavras for menor que ~35px, junta no mesmo texto/bloco
                if gap < 35:
                    current_cluster.append(w)
                else:
                    clusters.append(current_cluster)
                    current_cluster = [w]
        if current_cluster:
            clusters.append(current_cluster)

        # Para cada cluster da linha, calcula a posição em polegadas
        line_clusters = []
        for c in clusters:
            texto_cluster = " ".join(x["text"] for x in c)
            start_x = c[0]["x"]
            rel_x = start_x / float(img_width)
            pos_in = round(rel_x * printable_width_in, 2)
            line_clusters.append({
                "text": texto_cluster,
                "pos_in": pos_in,
                "rel_x": rel_x
            })

        resultado_linhas.append(line_clusters)

    return resultado_linhas


def _converter_pdf_escaneado_para_word_ocr(caminho_pdf, caminho_saida_docx, callback_progresso=None, idioma_ocr=None, total_paginas=None):
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_TAB_ALIGNMENT
    import winocr

    doc = Document()
    
    # Define margens de 0.5 polegadas (1.27 cm) para maximizar área de layout fiel
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    printable_width_in = 7.27  # Largura útil A4 com margens de 0.5 in

    pdf = pdfium.PdfDocument(caminho_pdf)
    if total_paginas is None:
        total_paginas = len(pdf)

    # Idioma OCR
    if not idioma_ocr or idioma_ocr == "Nenhum disponível":
        try:
            import winrt.windows.media.ocr as ocr_engine
            langs = [l.language_tag for l in ocr_engine.OcrEngine.available_recognizer_languages]
            for tag in ("pt-BR", "pt-PT", "pt", "en-US", "en"):
                for l in langs:
                    if l.lower() == tag.lower() or l.lower().startswith(tag.lower() + "-"):
                        idioma_ocr = l
                        break
                if idioma_ocr:
                    break
            if not idioma_ocr and langs:
                idioma_ocr = langs[0]
        except Exception:
            idioma_ocr = None

    for i, page in enumerate(pdf):
        if callback_progresso:
            prog = int(((i + 1) / total_paginas) * 90) + 5
            callback_progresso(prog, f"[OCR] Mapeando layout da página escaneada {i+1}/{total_paginas}...")

        scale = 200 / 72.0
        image_pil = page.render(scale=scale).to_pil()
        img_w, img_h = image_pil.size
        
        res = None
        if idioma_ocr:
            try:
                res = winocr.recognize_pil_sync(image_pil, lang=idioma_ocr)
            except Exception:
                pass

        if i > 0:
            doc.add_page_break()

        linhas_estruturadas = reconstruir_layout_ocr_estruturado(res, img_w, img_h, printable_width_in=printable_width_in)

        if linhas_estruturadas:
            for line_clusters in linhas_estruturadas:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.15

                for idx, c in enumerate(line_clusters):
                    pos_in = c["pos_in"]
                    texto = c["text"]

                    if idx == 0:
                        # Primeiro cluster da linha
                        if pos_in > 0.4:
                            p.paragraph_format.tab_stops.add_tab_stop(Inches(pos_in))
                            p.add_run(f"\t{texto}")
                        else:
                            p.add_run(texto)
                    else:
                        # Clusters subsequentes: aplica Tab Stop na posição X relativa
                        if c["rel_x"] > 0.85:
                            p.paragraph_format.tab_stops.add_tab_stop(Inches(printable_width_in), alignment=WD_TAB_ALIGNMENT.RIGHT)
                        else:
                            p.paragraph_format.tab_stops.add_tab_stop(Inches(pos_in))
                        p.add_run(f"\t{texto}")
        else:
            p = doc.add_paragraph(f"[Página {i+1} - Sem texto extraível via OCR]")
            p.paragraph_format.space_after = Pt(2)

        page.close()

    pdf.close()
    return salvar_documento_word_seguro(doc, caminho_saida_docx, callback_progresso)


def resolver_caminho_saida_disponivel(caminho_saida_docx):
    """
    Verifica se o arquivo de destino está aberto em outro programa (Word) e retorna um caminho livre.
    """
    if not os.path.exists(caminho_saida_docx):
        return caminho_saida_docx
    try:
        with open(caminho_saida_docx, "a+b"):
            pass
        return caminho_saida_docx
    except (PermissionError, OSError):
        pasta, nome_completo = os.path.split(caminho_saida_docx)
        nome, ext = os.path.splitext(nome_completo)
        for i in range(1, 100):
            novo_caminho = os.path.join(pasta, f"{nome}_{i}{ext}")
            if not os.path.exists(novo_caminho):
                return novo_caminho
            try:
                with open(novo_caminho, "a+b"):
                    pass
                return novo_caminho
            except (PermissionError, OSError):
                continue
        return caminho_saida_docx


def salvar_documento_word_seguro(doc, caminho_saida_docx, callback_progresso=None):
    """
    Salva o documento Word de forma segura, criando uma nova versão numerada se o arquivo original estiver aberto.
    """
    try:
        doc.save(caminho_saida_docx)
        if callback_progresso:
            callback_progresso(100, "Conversão concluída com layout fiel preservado via OCR!")
        return caminho_saida_docx
    except (PermissionError, OSError):
        pasta, nome_completo = os.path.split(caminho_saida_docx)
        nome, ext = os.path.splitext(nome_completo)
        for i in range(1, 100):
            novo_caminho = os.path.join(pasta, f"{nome}_{i}{ext}")
            try:
                doc.save(novo_caminho)
                if callback_progresso:
                    callback_progresso(100, f"[AVISO] O arquivo '{nome_completo}' estava aberto no Word. Salvo como: {os.path.basename(novo_caminho)}")
                return novo_caminho
            except (PermissionError, OSError):
                continue
        raise PermissionError(f"Não foi possível salvar o arquivo Word. Por favor, feche o documento no Microsoft Word e tente novamente: {caminho_saida_docx}")


def converter_pdf_para_word(caminho_pdf, caminho_saida_docx=None, callback_progresso=None, idioma_ocr=None):
    """
    Converte um arquivo PDF para Word (.docx) mantendo o layout de textos e tabelas.
    Se o PDF for escaneado (imagem sem camada de texto), utiliza OCR para extrair texto editável.
    """
    if not os.path.exists(caminho_pdf):
        raise FileNotFoundError(f"Arquivo PDF não encontrado: {caminho_pdf}")

    if not caminho_saida_docx:
        caminho_saida_docx = os.path.splitext(caminho_pdf)[0] + ".docx"

    caminho_saida_docx = resolver_caminho_saida_disponivel(caminho_saida_docx)

    # Verificação prévia: se o PDF possui pouca ou nenhuma camada de texto digital (PDF escaneado / imagem)
    eh_escaneado = False
    total_paginas = 0
    try:
        pdf_temp = pdfium.PdfDocument(caminho_pdf)
        total_paginas = len(pdf_temp)
        texto_digital = ""
        for p in pdf_temp:
            tp = p.get_textpage()
            texto_digital += tp.get_text_range().strip()
            tp.close()
            p.close()
        pdf_temp.close()
        
        # Se a média de caracteres por página for menor que 15, trata como PDF escaneado (imagem)
        if total_paginas > 0 and (len(texto_digital) / total_paginas) < 15:
            eh_escaneado = True
    except Exception:
        pass

    if eh_escaneado:
        if callback_progresso:
            callback_progresso(10, "[OCR] PDF Escaneado / Imagem detectado! Extraindo texto editável via OCR do Windows...")
        return _converter_pdf_escaneado_para_word_ocr(caminho_pdf, caminho_saida_docx, callback_progresso, idioma_ocr, total_paginas)

    # Para PDFs digitais com texto vetorial
    try:
        from pdf2docx import Converter
        if callback_progresso:
            callback_progresso(10, "Iniciando conversão de PDF para Word...")
            
        cv = Converter(caminho_pdf)
        cv.convert(caminho_saida_docx, start=0, end=None)
        cv.close()

        if callback_progresso:
            callback_progresso(100, "Conversão concluída!")

        return caminho_saida_docx
    except Exception as exc:
        if callback_progresso:
            callback_progresso(20, f"[AVISO] pdf2docx indisponível ou falhou. Tentando extração via OCR...")
        return _converter_pdf_escaneado_para_word_ocr(caminho_pdf, caminho_saida_docx, callback_progresso, idioma_ocr)


def converter_word_para_pdf(caminho_word, caminho_saida_pdf=None, callback_progresso=None):
    """
    Converte um arquivo Word (.docx/.doc) para PDF.
    """
    if not os.path.exists(caminho_word):
        raise FileNotFoundError(f"Arquivo Word não encontrado: {caminho_word}")

    if not caminho_saida_pdf:
        caminho_saida_pdf = os.path.splitext(caminho_word)[0] + ".pdf"

    caminho_word_abs = os.path.abspath(caminho_word)
    caminho_saida_pdf_abs = os.path.abspath(caminho_saida_pdf)

    try:
        import win32com.client
        pythoncom = None
        try:
            import pythoncom
            pythoncom.CoInitialize()
        except Exception:
            pass

        if callback_progresso:
            callback_progresso(20, "Abrindo documento Word...")

        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False

        doc = word.Documents.Open(caminho_word_abs)
        
        if callback_progresso:
            callback_progresso(60, "Exportando para PDF...")

        # 17 = wdFormatPDF
        doc.SaveAs(caminho_saida_pdf_abs, FileFormat=17)
        doc.Close()
        word.Quit()

        if pythoncom:
            try:
                pythoncom.CoUninitialize()
            except Exception:
                pass

        if callback_progresso:
            callback_progresso(100, "Conversão de Word para PDF concluída!")

        return caminho_saida_pdf_abs
    except Exception as exc:
        # Fallback via docx2pdf
        try:
            from docx2pdf import convert
            if callback_progresso:
                callback_progresso(50, "Convertendo via docx2pdf...")
            convert(caminho_word_abs, caminho_saida_pdf_abs)
            if callback_progresso:
                callback_progresso(100, "Conversão concluída!")
            return caminho_saida_pdf_abs
        except Exception as inner_exc:
            raise RuntimeError(f"Erro ao converter Word para PDF: {exc} | {inner_exc}")


def juntar_pdfs(lista_pdfs, caminho_saida, callback_progresso=None):
    """
    Junta múltiplos arquivos PDF em um único PDF.
    """
    if not lista_pdfs:
        raise ValueError("Nenhum arquivo PDF fornecido para junção.")

    writer = PdfWriter()
    total_arquivos = len(lista_pdfs)

    for idx, arq in enumerate(lista_pdfs):
        if not os.path.exists(arq):
            continue
        reader = PdfReader(arq)
        for page in reader.pages:
            writer.add_page(page)

        if callback_progresso:
            prog = int(((idx + 1) / total_arquivos) * 100)
            callback_progresso(prog, f"Unindo arquivo {idx+1}/{total_arquivos}...")

    with open(caminho_saida, "wb") as f_out:
        writer.write(f_out)

    if callback_progresso:
        callback_progresso(100, "Junção de PDFs concluída!")

    return caminho_saida


def extrair_imagens_pdf(caminho_pdf, pasta_saida, dpi=200, callback_progresso=None):
    """
    Exporta cada página do PDF como imagem PNG em alta resolução.
    """
    if not os.path.exists(caminho_pdf):
        raise FileNotFoundError(f"Arquivo PDF não encontrado: {caminho_pdf}")

    os.makedirs(pasta_saida, exist_ok=True)
    pdf = pdfium.PdfDocument(caminho_pdf)
    total_paginas = len(pdf)
    imagens_geradas = []

    scale = dpi / 72.0

    for i, page in enumerate(pdf):
        image = page.render(scale=scale).to_pil()
        nome_img = f"pagina_{i+1:03d}.png"
        caminho_img = os.path.join(pasta_saida, nome_img)
        image.save(caminho_img)
        imagens_geradas.append(caminho_img)

        if callback_progresso:
            prog = int(((i + 1) / total_paginas) * 100)
            callback_progresso(prog, f"Exportando página {i+1}/{total_paginas} como imagem...")

        page.close()

    pdf.close()
    return imagens_geradas
